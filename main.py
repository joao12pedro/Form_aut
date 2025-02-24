import docx

# Cria um novo documento
doc = docx.Document()

# Adiciona um título ao documento
doc.add_heading('Meu Documento Word Simples', 0)

# Adiciona um parágrafo de texto
doc.add_paragraph('Este é um parágrafo simples no documento Word.')
doc.add_picture('teste.jpeg', width=docx.shared.Inches(2))

# Salva o documento no diretório atual
doc.save('documento_simples.docx')

print('Documento Word criado com sucesso!')
